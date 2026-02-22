"""Servico principal de geracao, normalizacao e exportacao do escopo do projeto."""

from pathlib import Path
import re
from copy import deepcopy
import logging

from app.core.constants import (
    DEFAULT_PROMPTS_DIR,
    DEFAULT_SCOPE_DIR,
    DEFAULT_SCOPE_DOCX_PATH,
    DEFAULT_SCOPE_NORMALIZED_TXT_PATH,
    DEFAULT_TEXT_MODEL,
    DEFAULT_WORD_NORMALIZATION_PROMPT_PATH,
    DEFAULT_WORD_TEMPLATE_PATH,
)
from app.core.exceptions import (
    DependencyNotInstalledError,
    DocumentGenerationError,
    PromptDirectoryNotFoundError,
    PromptFileNotFoundError,
    ScopeGenerationError,
    SectionFileNotFoundError,
)
from app.services.ai.generative_service import GenerativeService

logger = logging.getLogger(__name__)
INVISIBLE_CHARS_RE = re.compile(r"[\u200b\u200c\u200d\u200e\u200f\u2060\ufeff]")
CODIGO_REQUISITO_RE = re.compile(r"^(RF|RNF|RB)-\d+")
SUBTITULO_BLOCO_RE = re.compile(
    r"^(Regras de Neg[oó]cio|Regras / Crit[eé]rios|Impacta:)$",
    re.IGNORECASE,
)


class ProjectScopeService:
    """Servico responsavel por gerar, normalizar e exportar o escopo do projeto."""

    def __init__(
        self,
        prompts_dir: str = str(DEFAULT_PROMPTS_DIR),
        output_dir: str = str(DEFAULT_SCOPE_DIR),
        prompt_normalizacao_path: str = str(DEFAULT_WORD_NORMALIZATION_PROMPT_PATH),
    ):
        """Inicializa caminhos de entrada/saida e o cliente de geracao de texto."""
        self.prompts_dir = Path(prompts_dir)
        self.output_dir = Path(output_dir)
        self.prompt_normalizacao_path = Path(prompt_normalizacao_path)
        self.ai_service = GenerativeService()

    def gerar_documento_word_de_txts(
        self,
        arquivos_txt: list[str],
        arquivo_saida: str = str(DEFAULT_SCOPE_DOCX_PATH),
        arquivo_texto_normalizado_saida: str = str(DEFAULT_SCOPE_NORMALIZED_TXT_PATH),
        normalizar_com_ia: bool = False,
        model_normalizacao: str = DEFAULT_TEXT_MODEL,
    ) -> str:
        """Gera DOCX a partir de secoes txt ja existentes, com normalizacao opcional."""
        if not arquivos_txt:
            raise ScopeGenerationError("Nenhum arquivo .txt informado para gerar o Word.")

        secoes_geradas: list[tuple[int, str]] = []
        for arquivo in arquivos_txt:
            path = Path(arquivo)
            if not path.exists():
                raise SectionFileNotFoundError(f"Arquivo de secao nao encontrado: {path}")

            conteudo = self._limpar_blocos_markdown(path.read_text(encoding="utf-8"))
            secoes_geradas.append((self._ordem_secao_saida(path.name), conteudo))

        secoes_ordenadas = sorted(secoes_geradas, key=lambda item: item[0])
        if normalizar_com_ia:
            # Normaliza o texto consolidado apenas uma vez para evitar divergencia entre secoes.
            logger.info(
                "Normalizando texto consolidado de %d secoes antes de gerar o Word",
                len(secoes_ordenadas),
            )
            texto_consolidado = self._consolidar_secoes_para_normalizacao(secoes_ordenadas)
            texto_normalizado = self._normalizar_texto_para_word_com_ia(
                texto=texto_consolidado,
                model=model_normalizacao,
            )
            self._salvar_texto_normalizado(
                texto=texto_normalizado,
                arquivo_saida=arquivo_texto_normalizado_saida,
            )
            secoes_ordenadas = [(secoes_ordenadas[0][0], texto_normalizado)]

        logger.info("Gerando Word a partir de %d arquivos txt", len(secoes_ordenadas))
        return self._gerar_documento_word(secoes_ordenadas, arquivo_saida=arquivo_saida)

    def gerar_documentos(
        self,
        transcricao: str,
        model: str = DEFAULT_TEXT_MODEL,
        model_normalizacao: str = DEFAULT_TEXT_MODEL,
        arquivo_texto_normalizado_saida: str = str(DEFAULT_SCOPE_NORMALIZED_TXT_PATH),
    ) -> tuple[list[str], str]:
        """Gera secoes de escopo via prompts, normaliza o consolidado e cria o DOCX final."""
        logger.info("Carregando prompts de: %s", self.prompts_dir)
        prompts = self._listar_prompts_ordenados()
        self.output_dir.mkdir(parents=True, exist_ok=True)
        logger.info("Prompts encontrados: %d", len(prompts))

        arquivos_gerados: list[str] = []
        secoes_geradas: list[tuple[int, str]] = []
        respostas_por_numero: dict[int, str] = {}
        for prompt_path in prompts:
            # A ordem do prompt define dependencias entre secoes (ex.: prompt 4 e 5 usam prompt 2).
            numero_prompt = self._ordem_prompt(prompt_path)
            logger.info("Processando prompt %s (%s)", numero_prompt, prompt_path.name)
            prompt_template = prompt_path.read_text(encoding="utf-8")
            prompt_final = self._montar_prompt_final(
                prompt_template=prompt_template,
                numero_prompt=numero_prompt,
                transcricao=transcricao,
                respostas_por_numero=respostas_por_numero,
            )

            logger.info("Chamando IA para prompt %s", numero_prompt)
            resposta = self.ai_service.gerar_texto(prompt_final, model=model)
            resposta_limpa = self._limpar_blocos_markdown(resposta or "")
            respostas_por_numero[numero_prompt] = resposta_limpa

            nome_saida = self._nome_saida(prompt_path.name)
            arquivo_saida = self.output_dir / nome_saida
            arquivo_saida.write_text(resposta_limpa, encoding="utf-8")
            arquivos_gerados.append(str(arquivo_saida))
            secoes_geradas.append((numero_prompt, resposta_limpa))
            logger.info(
                "Resposta do prompt %s salva em %s (%d caracteres)",
                numero_prompt,
                arquivo_saida,
                len(resposta_limpa),
            )

        secoes_geradas = sorted(secoes_geradas, key=lambda item: item[0])
        logger.info("Normalizando texto consolidado final com %d secoes", len(secoes_geradas))
        texto_consolidado = self._consolidar_secoes_para_normalizacao(secoes_geradas)
        texto_normalizado = self._normalizar_texto_para_word_com_ia(
            texto=texto_consolidado,
            model=model_normalizacao,
        )
        self._salvar_texto_normalizado(
            texto=texto_normalizado,
            arquivo_saida=arquivo_texto_normalizado_saida,
        )

        logger.info("Gerando documento Word consolidado")
        caminho_docx = self._gerar_documento_word([(secoes_geradas[0][0], texto_normalizado)])
        logger.info("Documento Word final salvo em: %s", caminho_docx)
        return arquivos_gerados, caminho_docx

    @staticmethod
    def _salvar_texto_normalizado(texto: str, arquivo_saida: str) -> str:
        """Persiste o texto normalizado consolidado em disco."""
        saida = Path(arquivo_saida)
        saida.parent.mkdir(parents=True, exist_ok=True)
        saida.write_text((texto or "").strip(), encoding="utf-8")
        logger.info("Texto normalizado salvo em: %s", saida)
        return str(saida)

    @staticmethod
    def _consolidar_secoes_para_normalizacao(secoes_geradas: list[tuple[int, str]]) -> str:
        """Consolida as secoes em um unico texto para normalizacao de formato."""
        textos: list[str] = []
        for _, conteudo in secoes_geradas:
            texto = (conteudo or "").strip()
            if not texto:
                continue
            textos.append(texto)
        return "\n\n".join(textos).strip()

    def _listar_prompts_ordenados(self) -> list[Path]:
        """Lista e ordena os prompts validos por numero no nome do arquivo."""
        if not self.prompts_dir.exists():
            raise PromptDirectoryNotFoundError(
                f"Diretorio de prompts nao encontrado: {self.prompts_dir}"
            )

        arquivos = []
        for path in self.prompts_dir.glob("*.txt"):
            if not path.is_file():
                continue
            if not re.match(r"^\(\d+\)\s.*\.txt$", path.name):
                continue
            arquivos.append(path)

        arquivos = sorted(arquivos, key=self._ordem_prompt)
        if not arquivos:
            raise PromptFileNotFoundError(
                f"Nenhum prompt .txt encontrado em: {self.prompts_dir}"
            )
        return arquivos

    @staticmethod
    def _ordem_prompt(path: Path) -> int:
        """Extrai o numero do prompt para ordenacao deterministica."""
        match = re.search(r"\((\d+)\)", path.name)
        if match:
            return int(match.group(1))
        return 10**9

    @staticmethod
    def _nome_saida(nome_arquivo_prompt: str) -> str:
        """Converte nome de prompt em nome de arquivo de secao padronizado."""
        numero_match = re.search(r"\((\d+)\)", nome_arquivo_prompt)
        numero = int(numero_match.group(1)) if numero_match else 0

        base = re.sub(r"^\(\d+\)\s*", "", nome_arquivo_prompt)
        base = base.replace("Prompt_", "")
        base = base.replace(".txt", "")
        base = base.strip().replace(" ", "-").lower()
        return f"{numero:02d}_{base}.txt"

    @staticmethod
    def _montar_prompt_final(
        prompt_template: str,
        numero_prompt: int,
        transcricao: str,
        respostas_por_numero: dict[int, str],
    ) -> str:
        """Monta prompt final com transcricao e dependencias de secoes anteriores."""
        prompt_final = prompt_template.replace("{transcricao}", transcricao)

        if numero_prompt in {4, 5}:
            requisitos_funcionais = respostas_por_numero.get(2, "").strip()
            if requisitos_funcionais:
                logger.info(
                    "Injetando requisitos funcionais do prompt 2 no prompt %s",
                    numero_prompt,
                )
                if "{requisitos_funcionais}" in prompt_final:
                    prompt_final = prompt_final.replace(
                        "{requisitos_funcionais}",
                        requisitos_funcionais,
                    )
                else:
                    prompt_final += (
                        "\n\nRequisitos Funcionais ja gerados (Prompt 2):\n"
                        f"{requisitos_funcionais}\n"
                    )

        return prompt_final

    def _normalizar_texto_para_word_com_ia(self, texto: str, model: str) -> str:
        """Normaliza o texto consolidado para melhor estrutura antes da exportacao Word."""
        texto_base = (texto or "").strip()
        if not texto_base:
            return texto_base

        try:
            template = self.prompt_normalizacao_path.read_text(encoding="utf-8")
        except FileNotFoundError:
            logger.warning(
                "Prompt de normalizacao nao encontrado em %s. Usando texto original.",
                self.prompt_normalizacao_path,
            )
            return texto_base

        prompt = template.replace("{texto}", texto_base)
        try:
            logger.info("Normalizando texto com IA para formato Word")
            resposta = self.ai_service.gerar_texto(prompt, model=model)
            normalizado = self._limpar_blocos_markdown(resposta or "")
            if not normalizado:
                return texto_base
            return self._aplicar_espacamento_semantico(normalizado)
        except Exception as exc:
            logger.warning("Falha na normalizacao com IA. Usando texto original. Erro: %s", exc)
            return self._aplicar_espacamento_semantico(texto_base)

    def _gerar_documento_word(
        self,
        secoes_geradas: list[tuple[int, str]],
        modelo_path: str = str(DEFAULT_WORD_TEMPLATE_PATH),
        arquivo_saida: str = str(DEFAULT_SCOPE_DOCX_PATH),
    ) -> str:
        """Monta o DOCX final usando template base e secoes ordenadas."""
        try:
            from docx import Document
        except ModuleNotFoundError as exc:
            raise DependencyNotInstalledError(
                "Dependencia 'python-docx' nao encontrada. Instale com: pip install python-docx"
            ) from exc

        if not secoes_geradas:
            raise DocumentGenerationError("Nenhuma secao de escopo foi gerada para exportacao Word.")

        modelo = Path(modelo_path)
        if not modelo.exists():
            raise DocumentGenerationError(f"Modelo Word nao encontrado: {modelo}")

        logger.info("Montando Word com modelo: %s", modelo)
        documento = Document(str(modelo))
        if not documento.paragraphs:
            raise DocumentGenerationError(
                "O modelo Word nao possui paragrafo base para replicacao."
            )

        paragrafo_base = deepcopy(documento.paragraphs[0]._p)
        self._remover_paragrafos_extras(documento)

        for indice, (_, conteudo) in enumerate(secoes_geradas):
            if indice > 0:
                # Mantem separacao visual entre secoes no documento final.
                documento.add_page_break()
                documento._body._element.append(deepcopy(paragrafo_base))
                logger.info("Nova pagina adicionada para secao %d", indice + 1)
            self._adicionar_texto(documento, conteudo)

        self._remover_paragrafos_vazios(documento)
        saida = Path(arquivo_saida)
        saida.parent.mkdir(parents=True, exist_ok=True)
        documento.save(str(saida))
        return str(saida)

    @staticmethod
    def _remover_paragrafos_extras(documento) -> None:
        """Remove paragrafos excedentes do template, preservando apenas a base."""
        paragrafos = list(documento.paragraphs)
        for paragrafo in paragrafos[1:]:
            elemento = paragrafo._element
            elemento.getparent().remove(elemento)

    @staticmethod
    def _adicionar_texto(documento, conteudo: str) -> None:
        """Interpreta marcacoes simples e adiciona o conteudo no documento Word."""
        texto = (conteudo or "").strip()
        if not texto:
            paragrafo = documento.add_paragraph("Sem conteudo gerado para esta secao.")
            ProjectScopeService._estilizar_paragrafo(paragrafo, tipo="normal")
            return

        linhas = ProjectScopeService._normalizar_linhas_para_word(texto)
        for linha_limpa in linhas:
            if not linha_limpa:
                # Linha vazia real para manter o mesmo efeito visual do TXT no DOCX.
                paragrafo_vazio = documento.add_paragraph("")
                ProjectScopeService._estilizar_paragrafo(paragrafo_vazio, tipo="normal")
                continue

            inicio_de_bloco = CODIGO_REQUISITO_RE.match(linha_limpa) is not None
            subtitulo_bloco = SUBTITULO_BLOCO_RE.match(linha_limpa) is not None
            # Mantido explicito para facilitar ajustes finos de layout no futuro.
            if inicio_de_bloco:
                espaco_antes = 0
            elif subtitulo_bloco:
                espaco_antes = 0
            else:
                espaco_antes = 0

            if linha_limpa.startswith("## "):
                titulo = linha_limpa[3:].strip()
                ProjectScopeService._adicionar_titulo(documento, titulo, level=2)
                continue

            if linha_limpa.startswith("### "):
                titulo = linha_limpa[4:].strip()
                ProjectScopeService._adicionar_titulo(documento, titulo, level=3)
                continue

            if linha_limpa.startswith("- ") or linha_limpa.startswith("\u2022 "):
                paragrafo = ProjectScopeService._adicionar_bullet(documento)
                ProjectScopeService._adicionar_texto_com_negrito(paragrafo, linha_limpa[2:].strip())
                ProjectScopeService._estilizar_paragrafo(
                    paragrafo,
                    tipo="lista",
                    space_before_pt=espaco_antes,
                )
                continue

            paragrafo = documento.add_paragraph()
            ProjectScopeService._adicionar_texto_com_negrito(paragrafo, linha_limpa)
            ProjectScopeService._estilizar_paragrafo(
                paragrafo,
                tipo="normal",
                space_before_pt=espaco_antes,
            )

    @staticmethod
    def _limpar_blocos_markdown(texto: str) -> str:
        """Remove cercas markdown e reduz ruido de quebras de linha."""
        if not texto:
            return ""

        linhas = texto.splitlines()
        linhas_filtradas: list[str] = []
        vazio_anterior = False
        for linha in linhas:
            linha_limpa = ProjectScopeService._sanitizar_linha(linha)
            linha_stripped = linha_limpa.lower()
            if linha_stripped.startswith("```"):
                continue
            if not linha_limpa:
                if vazio_anterior:
                    continue
                vazio_anterior = True
                linhas_filtradas.append("")
                continue

            vazio_anterior = False
            linhas_filtradas.append(linha_limpa)
        return "\n".join(linhas_filtradas).strip()

    @staticmethod
    def _adicionar_texto_com_negrito(paragrafo, texto: str) -> None:
        """Converte marcacao **texto** em runs com negrito no Word."""
        partes = re.split(r"(\*\*.*?\*\*)", texto)
        for parte in partes:
            if not parte:
                continue
            if parte.startswith("**") and parte.endswith("**") and len(parte) >= 4:
                run = paragrafo.add_run(parte[2:-2])
                run.bold = True
                ProjectScopeService._estilizar_run(run)
            else:
                run = paragrafo.add_run(parte)
                ProjectScopeService._estilizar_run(run)

    @staticmethod
    def _adicionar_bullet(documento):
        """Adiciona paragrafo de lista com fallback para modelos sem estilo de bullet."""
        for estilo in ("List Bullet", "Lista com marcadores"):
            try:
                return documento.add_paragraph(style=estilo)
            except KeyError:
                continue

        paragrafo = documento.add_paragraph()
        run = paragrafo.add_run("- ")
        ProjectScopeService._estilizar_run(run)
        return paragrafo

    @staticmethod
    def _adicionar_titulo(documento, titulo: str, level: int) -> None:
        """Adiciona titulo com fallback para paragrafos quando estilo nao existe."""
        paragrafo = None
        try:
            paragrafo = documento.add_heading(titulo, level=level)
        except KeyError:
            paragrafo = documento.add_paragraph()
            run = paragrafo.add_run(titulo)
            run.bold = True

        ProjectScopeService._estilizar_paragrafo(paragrafo, tipo="titulo")
        for run in paragrafo.runs:
            run.bold = True
            ProjectScopeService._estilizar_run(run)

    @staticmethod
    def _estilizar_paragrafo(paragrafo, tipo: str, space_before_pt: int = 0) -> None:
        """Aplica espacamento padrao para titulos, listas e paragrafos comuns."""
        try:
            from docx.shared import Pt
        except ModuleNotFoundError:
            return

        formato = paragrafo.paragraph_format
        formato.space_before = Pt(space_before_pt)
        if tipo == "titulo":
            formato.space_after = Pt(8)
            formato.line_spacing = 1.0
            return
        if tipo == "lista":
            formato.space_after = Pt(2)
            formato.line_spacing = 1.0
            return

        formato.space_after = Pt(3)
        formato.line_spacing = 1.0

    @staticmethod
    def _estilizar_run(run) -> None:
        """Forca tipografia consistente para qualquer trecho inserido no DOCX."""
        try:
            from docx.shared import RGBColor
        except ModuleNotFoundError:
            return

        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(0, 0, 0)

    @staticmethod
    def _ordem_secao_saida(nome_arquivo: str) -> int:
        """Extrai o prefixo numerico de arquivos de secao para ordenacao."""
        match = re.match(r"^(\d+)_", nome_arquivo)
        if match:
            return int(match.group(1))
        return 10**9

    @staticmethod
    def _normalizar_linhas_para_word(texto: str) -> list[str]:
        """Normaliza linhas para exportacao, preservando no maximo uma linha vazia sequencial."""
        linhas_saida: list[str] = []
        vazio_anterior = False
        for linha in texto.splitlines():
            linha_limpa = ProjectScopeService._sanitizar_linha(linha)
            if not linha_limpa:
                if vazio_anterior:
                    continue
                linhas_saida.append("")
                vazio_anterior = True
                continue
            vazio_anterior = False
            if re.fullmatch(r"_+", linha_limpa):
                continue
            linhas_saida.append(linha_limpa)
        return linhas_saida

    @staticmethod
    def _aplicar_espacamento_semantico(texto: str) -> str:
        """Insere quebras entre blocos logicos (RF/RNF/RB e subtitulos)."""
        linhas = texto.splitlines()
        saida: list[str] = []
        ultimo_nao_vazio = ""

        for linha in linhas:
            atual = ProjectScopeService._sanitizar_linha(linha)
            if not atual:
                if saida and saida[-1] != "":
                    saida.append("")
                continue

            novo_bloco = (
                CODIGO_REQUISITO_RE.match(atual) is not None
                or SUBTITULO_BLOCO_RE.match(atual) is not None
            )
            if novo_bloco and ultimo_nao_vazio:
                if saida and saida[-1] != "":
                    saida.append("")

            saida.append(atual)
            ultimo_nao_vazio = atual

        while saida and not saida[-1]:
            saida.pop()
        return "\n".join(saida)

    @staticmethod
    def _sanitizar_linha(linha: str) -> str:
        """Remove caracteres invisiveis e espacos especiais de uma linha."""
        linha_sem_invisivel = INVISIBLE_CHARS_RE.sub("", linha or "")
        linha_sem_invisivel = linha_sem_invisivel.replace("\u00a0", " ").replace("\u2003", " ")
        return linha_sem_invisivel.strip()

    @staticmethod
    def _remover_paragrafos_vazios(documento) -> None:
        """Remove vazios de borda sem eliminar linhas em branco entre blocos validos."""
        paragrafos = list(documento.paragraphs)
        for indice, paragrafo in enumerate(paragrafos):
            elemento = paragrafo._element
            if ProjectScopeService._elemento_contem_tag(elemento, "drawing"):
                continue
            if ProjectScopeService._elemento_contem_quebra_pagina(elemento):
                continue

            texto_limpo = INVISIBLE_CHARS_RE.sub("", (paragrafo.text or "")).strip()
            if texto_limpo:
                continue

            texto_anterior = ""
            if indice > 0:
                texto_anterior = INVISIBLE_CHARS_RE.sub(
                    "",
                    (paragrafos[indice - 1].text or ""),
                ).strip()

            texto_proximo = ""
            if indice < len(paragrafos) - 1:
                texto_proximo = INVISIBLE_CHARS_RE.sub(
                    "",
                    (paragrafos[indice + 1].text or ""),
                ).strip()

            if texto_anterior and texto_proximo:
                continue

            elemento.getparent().remove(elemento)

    @staticmethod
    def _elemento_contem_tag(elemento, tag_local: str) -> bool:
        """Verifica se o elemento XML contem determinada tag local."""
        for item in elemento.iter():
            if str(item.tag).endswith(f"}}{tag_local}"):
                return True
        return False

    @staticmethod
    def _elemento_contem_quebra_pagina(elemento) -> bool:
        """Verifica se o elemento contem quebra de pagina explicita."""
        for item in elemento.iter():
            if not str(item.tag).endswith("}br"):
                continue
            tipo = item.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type")
            if tipo == "page":
                return True
        return False
